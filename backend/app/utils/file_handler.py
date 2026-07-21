import os, hashlib, aiofiles, asyncio, sys
from langchain_core.documents import Document

from app.core.logger_handler import logger
from app.utils.path_tool import get_abstract_path
from langchain_community.document_loaders import PyPDFLoader, TextLoader, UnstructuredPDFLoader, UnstructuredMarkdownLoader, UnstructuredPowerPointLoader

class FontBBoxStreamFilter:
    """
    stderr 流过滤器——过滤掉 PyMuPDF 渲染时产生的无用 FontBBox 警告信息。

    PyMuPDF 在打开某些 PDF 时会在 stderr 输出大量形如
    "FontBBox from font descriptor" 的警告，这些警告对业务逻辑无意义，
    但会污染日志。通过替换 sys.stderr 为本过滤器，在写入前检查内容，
    仅放行不包含该关键字的输出。

    使用方式：实例化后赋值给 sys.stderr，之后所有 stderr 输出都会经过过滤。
    """

    def __init__(self, stream):
        """
        初始化过滤器。

        Args:
            stream: 原始的 stderr 流对象，过滤后的数据将写入此流。
        """
        self.stream = stream

    def write(self, data):
        """
        写入数据，过滤掉包含 FontBBox 警告的内容。

        Args:
            data: 要写入 stderr 的字符串数据。
        """
        if 'FontBBox from font descriptor' not in data:
            self.stream.write(data)

    def flush(self):
        """刷新底层流的缓冲区。"""
        self.stream.flush()

sys.stderr = FontBBoxStreamFilter(sys.stderr)

async def get_file_md5_hex(file_path: str) -> str:
    """
    异步计算文件的 MD5 值，用于文件去重和缓存键生成。

    采用分块读取（每次 1024 字节）的方式处理大文件，避免一次性加载全部内容到内存。
    如果文件不存在或读取失败，返回空字符串。

    Args:
        file_path: 文件路径，支持绝对路径和相对路径（相对路径会通过 get_abstract_path 转换）。

    Returns:
        文件内容的 MD5 十六进制摘要字符串（32位），计算失败时返回空字符串 ""。
    """
    # 处理路径，确保使用绝对路径
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    
    if not os.path.exists(abs_file_path):
        logger.error(f"【md5计算】文件路径 {abs_file_path} 不存在")
        return ""

    if not os.path.isfile(abs_file_path):
        logger.error(f"【md5计算】文件路径 {abs_file_path} 不是文件")
        return ""

    md5_object = hashlib.md5()
    chunk_size = 1024
    try:
        async with aiofiles.open(abs_file_path, "rb") as f:
            while chunk := await f.read(chunk_size):
                md5_object.update(chunk)
    except Exception as e:
        logger.error(f"【md5计算】读取文件 {abs_file_path} 时出错: {e}")
        return ""

    return md5_object.hexdigest()

async def listdir_allowed_type(path: str, allowed_types: tuple[str]) -> tuple:
    """
    异步获取指定目录下所有符合指定扩展名的文件路径。

    使用 asyncio.to_thread 将阻塞的 os.listdir 调用放入线程池执行，
    避免阻塞事件循环。仅返回文件（不包含子目录），且文件名必须以 allowed_types 中的后缀结尾。

    Args:
        path: 目录路径，支持绝对路径和相对路径。
        allowed_types: 允许的文件扩展名元组，例如 ('.pdf', '.txt', '.docx')。

    Returns:
        符合条件的文件绝对路径元组。目录不存在或不是目录时返回空元组 ()。
    """
    # 处理路径，确保使用绝对路径
    abs_path = get_abstract_path(path) if not os.path.isabs(path) else path
    
    if not os.path.exists(abs_path):
        logger.error(f"【文件列表】目录路径 {abs_path} 不存在")
        return ()

    if not os.path.isdir(abs_path):
        logger.error(f"【文件列表】目录路径 {abs_path} 不是目录")
        return ()

    file_list = []
    for f in await asyncio.to_thread(os.listdir, abs_path):
        if f.endswith(allowed_types):
            file_path = os.path.join(abs_path, f)
            file_list.append(file_path)

    return tuple(file_list)



async def pdf_loader(file_path: str, password: str = None) -> list[Document]:
    """
    异步加载 PDF 文件内容，返回 LangChain Document 对象列表。

    加载策略（降级机制）：
    1. 如果提供了密码，使用 PyPDFLoader 解密加载。
    2. 优先尝试 UnstructuredPDFLoader（支持提取图片旁的文字和混合内容）。
    3. 如果 UnstructuredPDFLoader 失败或提取结果为空，降级使用 PyPDFLoader。

    Args:
        file_path: PDF 文件路径，支持绝对路径和相对路径。
        password: PDF 文件的解密密码，如果文件未加密则传 None。

    Returns:
        LangChain Document 列表，每个 Document 对应 PDF 的一页，
        page_content 为该页的文本内容。加载失败时返回空列表。
    """
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    
    if password:
        loader = PyPDFLoader(abs_file_path, password=password)
        return await asyncio.to_thread(loader.load)
    
    try:
        loader = UnstructuredPDFLoader(abs_file_path)
        docs = await asyncio.to_thread(loader.load)
        if docs and any(len(doc.page_content.strip()) > 0 for doc in docs):
            return docs
    except Exception as e:
        logger.warning(f"【PDF加载】UnstructuredPDFLoader失败，尝试PyPDFLoader: {e}")
    
    loader = PyPDFLoader(abs_file_path)
    return await asyncio.to_thread(loader.load)


async def txt_loader(file_path: str) -> list[Document]:
    """
    异步加载 TXT 文本文件内容，返回 LangChain Document 对象列表。

    编码检测策略：依次尝试 utf-8 和 gbk 两种常见编码，
    因为中文环境下部分文本文件使用 gbk 编码保存。
    如果所有编码都失败，返回空列表。

    Args:
        file_path: TXT 文件路径，支持绝对路径和相对路径。

    Returns:
        LangChain Document 列表（通常只有一个元素，包含整个文件的文本内容）。
        加载失败时返回空列表。
    """
    # 处理路径，确保使用绝对路径
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    
    # 使用不同的编码加载文件
    encodings = ['utf-8', 'gbk']
    for encoding in encodings:
        try:
            loader = TextLoader(abs_file_path, encoding=encoding)
            return await asyncio.to_thread(loader.load)
        except Exception as e:
            logger.error(f"【文本文件加载】使用编码 {encoding} 加载文件 {abs_file_path} 时出错: {e}")
            continue
    # 所有编码都失败，返回空列表
    return []


def _load_docx(file_path: str) -> list[Document]:
    """使用 python-docx 提取 DOCX 段落和表格文本。"""
    from docx import Document as DocxDocument

    docx = DocxDocument(file_path)
    blocks: list[str] = []

    for paragraph in docx.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in docx.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append("\t".join(cells))

    content = "\n\n".join(blocks).strip()
    if not content:
        return []
    return [Document(page_content=content, metadata={"source": file_path})]


async def word_loader(file_path: str) -> list[Document]:
    """异步加载 DOCX 文件内容，返回 LangChain Document 对象列表。"""
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    try:
        return await asyncio.to_thread(_load_docx, abs_file_path)
    except Exception as e:
        logger.error(f"【WORD文件加载】加载文件 {abs_file_path} 时出错: {e}", exc_info=True)
        return []
def _load_markdown_as_text(abs_file_path: str) -> list[Document]:
    """Load Markdown through TextLoader when optional parsers are unavailable."""
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return TextLoader(abs_file_path, encoding=encoding).load()
        except Exception:
            continue
    return []


async def markdown_loader(file_path: str) -> list[Document]:
    """Load Markdown asynchronously, falling back to TextLoader if necessary."""
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    try:
        loader = UnstructuredMarkdownLoader(abs_file_path, mode="single")
        return await asyncio.to_thread(loader.load)
    except Exception as e:
        logger.warning(f"[Markdown loader] Unstructured failed; using TextLoader fallback: {e}")
        return await asyncio.to_thread(_load_markdown_as_text, abs_file_path)


async def ppt_loader(file_path: str) -> list[Document]:
    """
    异步加载 PPT/PPTX 演示文稿内容，返回 LangChain Document 对象列表。

    使用 UnstructuredPowerPointLoader 以 "single" 模式加载，
    将整个 PPT 文件的文本内容作为一个 Document 返回。

    Args:
        file_path: PPT 文件路径（.ppt 或 .pptx），支持绝对路径和相对路径。

    Returns:
        LangChain Document 列表。加载失败时返回空列表。
    """
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    try:
        loader = UnstructuredPowerPointLoader(abs_file_path, mode="single")
        return await asyncio.to_thread(loader.load)
    except Exception as e:
        logger.error(f"【PPT文件加载】加载文件 {abs_file_path} 时出错: {e}")
        return []


def get_file_md5_hex_sync(file_path: str) -> str:
    """
    同步计算文件的 MD5 值，用于多线程（ThreadPoolExecutor）环境。

    与异步版本 get_file_md5_hex 功能相同，但使用同步文件 I/O，
    适用于线程池中执行的任务（如 SSE 上传流程中的 _sync_slice_file）。

    Args:
        file_path: 文件路径，支持绝对路径和相对路径。

    Returns:
        文件内容的 MD5 十六进制摘要字符串（32位），计算失败时返回空字符串 ""。
    """
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    
    if not os.path.exists(abs_file_path):
        logger.error(f"【md5计算】文件路径 {abs_file_path} 不存在")
        return ""

    if not os.path.isfile(abs_file_path):
        logger.error(f"【md5计算】文件路径 {abs_file_path} 不是文件")
        return ""

    md5_object = hashlib.md5()
    chunk_size = 1024
    try:
        with open(abs_file_path, "rb") as f:
            while chunk := f.read(chunk_size):
                md5_object.update(chunk)
    except Exception as e:
        logger.error(f"【md5计算】读取文件 {abs_file_path} 时出错: {e}")
        return ""

    return md5_object.hexdigest()


def pdf_loader_sync(file_path: str, password: str = None) -> list[Document]:
    """
    同步加载 PDF 文件内容，用于多线程（ThreadPoolExecutor）环境。

    加载策略与异步版 pdf_loader 完全一致（降级机制）：
    1. 有密码时使用 PyPDFLoader 解密加载。
    2. 优先尝试 UnstructuredPDFLoader。
    3. 失败后降级使用 PyPDFLoader。

    Args:
        file_path: PDF 文件路径，支持绝对路径和相对路径。
        password: PDF 文件的解密密码，未加密时传 None。

    Returns:
        LangChain Document 列表，每个 Document 对应 PDF 的一页。
        加载失败时返回空列表。
    """
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    
    if password:
        loader = PyPDFLoader(abs_file_path, password=password)
        return loader.load()
    
    try:
        loader = UnstructuredPDFLoader(abs_file_path)
        docs = loader.load()
        if docs and any(len(doc.page_content.strip()) > 0 for doc in docs):
            return docs
    except Exception as e:
        logger.warning(f"【PDF加载】UnstructuredPDFLoader失败，尝试PyPDFLoader: {e}")
    
    loader = PyPDFLoader(abs_file_path)
    return loader.load()


def txt_loader_sync(file_path: str) -> list[Document]:
    """
    同步加载 TXT 文本文件内容，用于多线程环境。

    编码检测策略与异步版一致：依次尝试 utf-8 和 gbk 两种编码。

    Args:
        file_path: TXT 文件路径，支持绝对路径和相对路径。

    Returns:
        LangChain Document 列表。加载失败时返回空列表。
    """
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    
    encodings = ['utf-8', 'gbk']
    for encoding in encodings:
        try:
            loader = TextLoader(abs_file_path, encoding=encoding)
            return loader.load()
        except Exception as e:
            logger.error(f"【文本文件加载】使用编码 {encoding} 加载文件 {abs_file_path} 时出错: {e}")
            continue
    return []


def word_loader_sync(file_path: str) -> list[Document]:
    """同步加载 DOCX 文件内容，用于多线程切片流程。"""
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    try:
        return _load_docx(abs_file_path)
    except Exception as e:
        logger.error(f"【WORD文件加载】加载文件 {abs_file_path} 时出错: {e}", exc_info=True)
        return []

def markdown_loader_sync(file_path: str) -> list[Document]:
    """Load Markdown synchronously, falling back to TextLoader if necessary."""
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    try:
        loader = UnstructuredMarkdownLoader(abs_file_path, mode="single")
        return loader.load()
    except Exception as e:
        logger.warning(f"[Markdown loader] Unstructured failed; using TextLoader fallback: {e}")
        return _load_markdown_as_text(abs_file_path)


def ppt_loader_sync(file_path: str) -> list[Document]:
    """
    同步加载 PPT/PPTX 演示文稿内容，用于多线程环境。

    使用 UnstructuredPowerPointLoader 以 "single" 模式加载。

    Args:
        file_path: PPT 文件路径（.ppt 或 .pptx），支持绝对路径和相对路径。

    Returns:
        LangChain Document 列表。加载失败时返回空列表。
    """
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    try:
        loader = UnstructuredPowerPointLoader(abs_file_path, mode="single")
        return loader.load()
    except Exception as e:
        logger.error(f"【PPT文件加载】加载文件 {abs_file_path} 时出错: {e}")
        return []
