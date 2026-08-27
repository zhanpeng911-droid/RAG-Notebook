"""file_handler 单元测试 —— MD5 计算、类型过滤、各格式加载与降级链。

langchain 加载器类在模块命名空间整体替换为可控桩，不触发真实解析；
DOCX 走真实 python-docx 生成临时文件验证段落/表格提取。
"""
import hashlib
import os

import pytest

from tests.helpers.unmock import LANGCHAIN_STACK, restore_real

restore_real(*LANGCHAIN_STACK)

# conftest 预注册的 app.utils.file_handler 假条目会短路父包绑定
restore_real("app.utils.file_handler")

import app  # noqa: E402,F401
import app.utils  # noqa: E402,F401
import app.utils.file_handler as fh  # noqa: E402
from app.utils.file_handler import (  # noqa: E402
    FontBBoxStreamFilter,
    get_file_md5_hex,
    get_file_md5_hex_sync,
    listdir_allowed_type,
)


def make_loader_cls(docs=None, error=None):
    class CtlLoader:
        calls = []

        def __init__(self, path, encoding=None, password=None, mode=None):
            type(self).calls.append({"path": path, "encoding": encoding,
                                     "password": password, "mode": mode})

        def load(self):
            if error:
                raise error
            return list(docs or [])

    return CtlLoader


# ---------- FontBBox 流过滤器 ----------

def test_fontbbox_filter_blocks_warning_and_passes_rest():
    class _Sink:
        def __init__(self):
            self.parts = []
            self.flushed = False

        def write(self, data):
            self.parts.append(data)

        def flush(self):
            self.flushed = True

    sink = _Sink()
    f = FontBBoxStreamFilter(sink)
    f.write("normal log\n")
    f.write("FontBBox from font descriptor xxx\n")
    f.flush()
    assert sink.parts == ["normal log\n"]
    assert sink.flushed is True


# ---------- MD5 ----------

@pytest.mark.asyncio
async def test_md5_async_matches_hashlib(tmp_path):
    f = tmp_path / "a.bin"
    data = b"x" * 3000  # 跨多个 1024 分块
    f.write_bytes(data)
    assert await get_file_md5_hex(str(f)) == hashlib.md5(data).hexdigest()


@pytest.mark.asyncio
async def test_md5_async_missing_and_dir_paths(tmp_path):
    assert await get_file_md5_hex(str(tmp_path / "ghost")) == ""
    assert await get_file_md5_hex(str(tmp_path)) == ""  # 目录不是文件


def test_md5_sync_happy_and_missing(tmp_path):
    f = tmp_path / "b.bin"
    f.write_bytes(b"payload")
    assert get_file_md5_hex_sync(str(f)) == hashlib.md5(b"payload").hexdigest()
    assert get_file_md5_hex_sync(str(tmp_path / "nope")) == ""


# ---------- 目录类型过滤 ----------

@pytest.mark.asyncio
async def test_listdir_allowed_type_filters(tmp_path):
    (tmp_path / "a.pdf").write_text("1")
    (tmp_path / "b.txt").write_text("2")
    (tmp_path / "c.exe").write_text("3")
    sub = tmp_path / "sub.pdf"
    sub.mkdir()

    out = await listdir_allowed_type(str(tmp_path), (".pdf", ".txt"))
    names = {os.path.basename(p) for p in out}
    # 现状实现仅按扩展名过滤，未做 isfile 校验（docstring 与实现不符，
    # 已登记遗留问题；R3 修复后此断言应改回不含 sub.pdf）
    assert names == {"a.pdf", "b.txt", "sub.pdf"}


@pytest.mark.asyncio
async def test_listdir_missing_or_file_returns_empty(tmp_path):
    assert await listdir_allowed_type(str(tmp_path / "none"), ()) == ()
    f = tmp_path / "f.txt"
    f.write_text("x")
    assert await listdir_allowed_type(str(f), (".txt",)) == ()


# ---------- PDF：密码优先 + Unstructured→PyPDF 降级 ----------

@pytest.mark.asyncio
async def test_pdf_loader_password_branch(monkeypatch, tmp_path):
    cls = make_loader_cls(docs=[_fake_doc("pwd page")])
    monkeypatch.setattr(fh, "PyPDFLoader", cls)
    out = await fh.pdf_loader(str(tmp_path / "p.pdf"), password="123")
    assert cls.calls[0]["password"] == "123"
    assert out[0].page_content == "pwd page"


@pytest.mark.asyncio
async def test_pdf_loader_unstructured_success_kept(monkeypatch, tmp_path):
    uni = make_loader_cls(docs=[_fake_doc("丰富内容")])
    py = make_loader_cls()
    monkeypatch.setattr(fh, "UnstructuredPDFLoader", uni)
    monkeypatch.setattr(fh, "PyPDFLoader", py)
    out = await fh.pdf_loader(str(tmp_path / "p.pdf"))
    assert out[0].page_content == "丰富内容"
    assert py.calls == []


@pytest.mark.asyncio
async def test_pdf_loader_unstructured_failure_falls_back(monkeypatch, tmp_path):
    uni = make_loader_cls(error=RuntimeError("依赖缺失"))
    py = make_loader_cls(docs=[_fake_doc("fallback")])
    monkeypatch.setattr(fh, "UnstructuredPDFLoader", uni)
    monkeypatch.setattr(fh, "PyPDFLoader", py)
    out = await fh.pdf_loader(str(tmp_path / "p.pdf"))
    assert out[0].page_content == "fallback"
    assert len(uni.calls) == 1


@pytest.mark.asyncio
async def test_pdf_loader_empty_unstructured_result_falls_back(monkeypatch, tmp_path):
    uni = make_loader_cls(docs=[_fake_doc("   ")])  # 只有空白内容视为失败
    py = make_loader_cls(docs=[_fake_doc("pypdf 页")])
    monkeypatch.setattr(fh, "UnstructuredPDFLoader", uni)
    monkeypatch.setattr(fh, "PyPDFLoader", py)
    out = await fh.pdf_loader(str(tmp_path / "p.pdf"))
    assert out[0].page_content == "pypdf 页"


def test_pdf_loader_sync_mirrors_degradation(monkeypatch, tmp_path):
    uni = make_loader_cls(error=IOError("sync broken"))
    py = make_loader_cls(docs=[_fake_doc("sync")])
    monkeypatch.setattr(fh, "UnstructuredPDFLoader", uni)
    monkeypatch.setattr(fh, "PyPDFLoader", py)
    out = fh.pdf_loader_sync(str(tmp_path / "p.pdf"))
    assert out[0].page_content == "sync"

    pw = make_loader_cls(docs=[_fake_doc("加密页")])
    monkeypatch.setattr(fh, "PyPDFLoader", pw)
    assert fh.pdf_loader_sync(str(tmp_path / "p.pdf"), password="k")[0] \
        .page_content == "加密页"


# ---------- TXT：编码降级 ----------

class ScriptedTextLoader:
    """utf-8 必失败、其余编码成功的桩：验证编码降级顺序。"""

    instances = []

    def __init__(self, path, encoding=None, password=None, mode=None):
        self.encoding = encoding
        type(self).instances.append(self)

    def load(self):
        if self.encoding == "utf-8":
            raise UnicodeDecodeError("utf-8", b"", 0, 1, "invalid")
        return [_fake_doc(f"decoded-{self.encoding}")]


@pytest.mark.asyncio
async def test_txt_loader_utf8_fails_then_gbk_succeeds(monkeypatch, tmp_path):
    ScriptedTextLoader.instances = []
    monkeypatch.setattr(fh, "TextLoader", ScriptedTextLoader)
    out = await fh.txt_loader(str(tmp_path / "x.txt"))
    assert out[0].page_content == "decoded-gbk"
    assert [i.encoding for i in ScriptedTextLoader.instances] == ["utf-8", "gbk"]

    ScriptedTextLoader.instances = []
    uni_fail = make_loader_cls(error=OSError("永不成功"))
    monkeypatch.setattr(fh, "TextLoader", uni_fail)
    assert await fh.txt_loader(str(tmp_path / "x.txt")) == []


def test_txt_loader_sync_encoding_chain(monkeypatch, tmp_path):
    ScriptedTextLoader.instances = []
    monkeypatch.setattr(fh, "TextLoader", ScriptedTextLoader)
    out = fh.txt_loader_sync(str(tmp_path / "s.txt"))
    assert out[0].page_content == "decoded-gbk"
    assert [i.encoding for i in ScriptedTextLoader.instances] == ["utf-8", "gbk"]

    bad = make_loader_cls(error=ValueError("no decode"))
    monkeypatch.setattr(fh, "TextLoader", bad)
    assert fh.txt_loader_sync(str(tmp_path / "s.txt")) == []


# ---------- DOCX：真实 python-docx ----------

@pytest.mark.asyncio
async def test_word_loader_extracts_paragraphs_and_tables(tmp_path):
    from docx import Document as DocxDoc
    d = DocxDoc()
    d.add_paragraph("标题段")
    d.add_paragraph("正文第二段")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "年龄"
    table.cell(1, 0).text = "张三"
    table.cell(1, 1).text = "21"
    fp = tmp_path / "t.docx"
    d.save(str(fp))

    out = await fh.word_loader(str(fp))
    assert len(out) == 1
    text = out[0].page_content
    assert "标题段" in text and "正文第二段" in text
    assert "姓名\t年龄" in text and "张三\t21" in text
    assert out[0].metadata["source"].endswith("t.docx")


@pytest.mark.asyncio
async def test_word_loader_blank_document_returns_empty(tmp_path):
    from docx import Document as DocxDoc
    d = DocxDoc()
    fp = tmp_path / "blank.docx"
    d.save(str(fp))
    assert await fh.word_loader(str(fp)) == []


@pytest.mark.asyncio
async def test_word_loader_broken_file_returns_empty(tmp_path):
    fp = tmp_path / "broken.docx"
    fp.write_bytes(b"not a zip")
    assert await fh.word_loader(str(fp)) == []
    assert fh.word_loader_sync(str(fp)) == []


# ---------- Markdown：Unstructured 失败回退 TextLoader 多编码 ----------

@pytest.mark.asyncio
async def test_markdown_loader_fallback_chain(monkeypatch, tmp_path):
    uni = make_loader_cls(error=ImportError("无可选依赖"))
    monkeypatch.setattr(fh, "UnstructuredMarkdownLoader", uni)

    ScriptedTextLoader.instances = []
    monkeypatch.setattr(fh, "TextLoader", ScriptedTextLoader)
    out = await fh.markdown_loader(str(tmp_path / "m.md"))
    assert out[0].page_content.startswith("decoded-")

    text_fail = make_loader_cls(error=OSError("全编码失败"))
    monkeypatch.setattr(fh, "TextLoader", text_fail)
    assert await fh.markdown_loader(str(tmp_path / "m.md")) == []


def test_markdown_loader_sync_fallback(monkeypatch, tmp_path):
    uni = make_loader_cls(error=ImportError("缺依赖"))
    monkeypatch.setattr(fh, "UnstructuredMarkdownLoader", uni)
    ScriptedTextLoader.instances = []
    monkeypatch.setattr(fh, "TextLoader", ScriptedTextLoader)
    out = fh.markdown_loader_sync(str(tmp_path / "ms.md"))
    assert out[0].page_content.startswith("decoded-")


# ---------- PPT ----------

@pytest.mark.asyncio
async def test_ppt_loader_success_and_error(monkeypatch, tmp_path):
    uni = make_loader_cls(docs=[_fake_doc("幻灯片文字")])
    monkeypatch.setattr(fh, "UnstructuredPowerPointLoader", uni)
    out = await fh.ppt_loader(str(tmp_path / "d.pptx"))
    assert out[0].page_content == "幻灯片文字"
    assert uni.calls[0]["mode"] == "single"

    err = make_loader_cls(error=RuntimeError("解析失败"))
    monkeypatch.setattr(fh, "UnstructuredPowerPointLoader", err)
    assert await fh.ppt_loader(str(tmp_path / "d.pptx")) == []
    assert fh.ppt_loader_sync(str(tmp_path / "d.pptx")) == []


# ---------- 工具 ----------

def _fake_doc(content):
    from langchain_core.documents import Document
    return Document(page_content=content, metadata={})
